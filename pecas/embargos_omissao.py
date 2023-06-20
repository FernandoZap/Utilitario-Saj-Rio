
# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Length
import docx

import os
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches,Cm
from . import funcoes_gerais,teses_embargos_omissao,estilos,funcoes_banco,funcoes_impressao
from .classes import Pasta



def peticoes2(modelo,teses,dados_compl,pasta):
    pass


def peticoes(modelo,teses,dados_compl,pasta):
    #lista = funcoes_banco.f001_sql(pasta)
    lista=funcoes_banco.view_pasta(pasta)
    pasta_saj = Pasta(lista[0],lista[1],lista[2],lista[3],lista[4],lista[5],lista[6],lista[7],lista[8],lista[9],lista[10],lista[11],lista[12],lista[13],lista[14],lista[15],lista[16],lista[17],lista[18])

    dados_pasta = funcoes_banco.f002_sql(pasta)

    doc = docx.Document()

    sec1 =  doc.sections[-1]
    ft1 = sec1.footer
    hd1 =  sec1.header

    ft1_pg =  ft1.add_paragraph('João Barbosa Advogado Associados')
    ft1_pg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    hd1_pg =  hd1.add_paragraph(pasta_saj.cod_cliente+' C3/'+pasta_saj.pasta+'/'+pasta_saj.cobertura)
    hd1_pg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    styles = doc.styles
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Times New Roman"
    p.font.size = Pt(13)
    #p.font.color.rgb=RGBColor(79, 129, 189)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #p.paragraph_format.first_line_indent = Inches(0.5)
    #p.paragraph_format.line_spacing = Inches(0.35)


    p2 = styles.add_style("Paragraph-2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Times New Roman"
    p2.font.size = Pt(12)
    #p2.font.color.rgb=RGBColor(79, 129, 189)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    #p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.left_indent = Inches(1.5)

    p3 = styles.add_style("Paragraph-3", WD_STYLE_TYPE.PARAGRAPH)
    p3.font.name = "Times New Roman"
    p3.font.size = Pt(12)
    p3.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    p4 = styles.add_style("Paragraph-4", WD_STYLE_TYPE.PARAGRAPH)
    p4.font.name = "Times New Roman"
    p4.font.size = Pt(13)
    p4.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p4.font.bold = True

    p41 = styles.add_style("Paragraph-41", WD_STYLE_TYPE.PARAGRAPH)
    p41.font.name = "Times New Roman"
    p41.font.size = Pt(12)
    p41.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p41.font.bold = True
    p41.font.italic = True


    # config inicial / logo / rodape
    #funcoes_impressao.pagina_inicial(doc,pasta_saj)
    '''
    for item in teses:
        if teses[item]=='S':
            imprime_tese(teses_embargos_omissao.teseConstrutor(dados_compl,dados_pasta,item),doc)
    '''
    imprime_tese1(pasta_saj,doc)

    # assinatura
    funcoes_impressao.pagina_de_encerramento(doc,pasta_saj)

    #nome_documento="c:/Avulsos/testes/embargos" +funcoes_gerais.data_doc() +".docx"
    nome_documento="/home/ubuntu/documentos/embargos" +funcoes_gerais.data_doc() +".docx"
    nome="embargos" + funcoes_gerais.data_doc() +".docx"
    doc.save(nome_documento)

    return nome


def imprime_tese2(tese,doc):
    pass


def imprime_tese1(pasta_saj,doc):

    #ft1_pg =  ft1.add_paragraph('*João Barbosa Advogado Associados*')
    #ft1_pg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    #hd1_pg =  hd1.add_paragraph(pasta_saj.cod_cliente+' C3/'+pasta_saj.pasta+'/'+pasta_saj.cobertura)
    #hd1_pg.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.RIGHT


    doc.add_paragraph('EXMO SR. DR. JUIZ DE DIREITO DO(A) '+pasta_saj.juizo+' DA COMARCA DE '+pasta_saj.comarca+'/'+pasta_saj.uf, style='Paragraph-4')
    doc.add_paragraph('',style='Paragraph-4')
    doc.add_paragraph('',style='Paragraph-4')
    doc.add_paragraph('Processo n. '+pasta_saj.nr_processo,style='Paragraph-41')


    para = doc.add_paragraph('',style='Paragraph')
    para.add_run(pasta_saj.cliente).bold = True
    para.add_run(', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da ')
    para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT').bold = True
    para.add_run(' que lhe promove ')
    para.add_run(pasta_saj.autor).bold = True
    para.add_run(', em trâmite perante este Douto Juízo, vem respeitosamente, à presença de V. Exa., ')
    para.add_run('requerer o DESARQUIVAMENTO, a fim de viabilizar a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e ofício único de pagamento.').bold=True








