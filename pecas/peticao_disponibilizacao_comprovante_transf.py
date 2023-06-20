from docx import Document
import docx

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches


#document = Document()
#styles =  document.styles

#Style Paragraph
#p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
#p.font.name = "Calibri"
#p.font.size = Pt(11)

#Style Heading 2222
'''
h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = styles["Heading 2"]
h2.font.name = "Calibri"
h2.font.size = Pt(13)
h2.font.color.rgb = RGBColor(78, 129, 189)
h2.font.bold = False

#Style Heading 3
h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
h3.base_style = styles["Heading 3"]
h3.font.name = "Calibri"
h3.font.size = Pt(12)
h3.font.color.rgb = RGBColor(78, 129, 189)
h3.font.bold = False
'''
def fun_peticao_02(p_pasta,p_autor,p_nr_processo,p_comarca,p_uf,p_cliente):
    doc = docx.Document()
    doc.add_heading('Peticao do documento', 0)
    para = doc.add_paragraph(
    '''GeeksforGeeks is a Computer Science portal for geeks.''')

    # Adding more content to paragraph and applying underline to them
    para.add_run(
    ''' It contains well written, well thought and well-explained ''').underline = True

    # Adding more content to paragraph
    para.add_run('''computer science and programming articles, quizzes etc.''')

    # Now save the document to a location

    doc.save('/home/ubuntu/documentos/teste2.docx')
    return 'teste2.docx'



def fun_peticao_01(p_pasta,p_cod_cliente,p_autor,p_nr_processo,p_comarca,p_uf,p_cliente,p_juizo):
    document = docx.Document()

    styles = document.styles
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    p.font.color.rgb=RGBColor(79, 129, 189)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)


    p2 = styles.add_style("Paragraph-2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Calibri"
    p2.font.size = Pt(11)
    p2.font.color.rgb=RGBColor(79, 129, 189)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.left_indent = Inches(1.5)

    document.add_heading('EXMO SR. DR. JUIZ DE DIREITO DO(A) '+p_juizo+' DA COMARCA DE '+p_comarca+'/'+p_uf, 2)
    document.add_heading('Processo: '+p_nr_processo)

    para= document.add_paragraph(p_cliente+', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da',style='Paragraph')
    under_and_bold = para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT')
    under_and_bold.underline=True
    under_and_bold.bold=True
    para.add_run(', que lhe promove ')
    para.add_run(p_autor).bold=True
    para.add_run(', em trâmite perante este Douto Juízo e Respectivo Cartório, vem, mui respeitosamente, à presença de V. Exa., informar para ao final requerer:')

    document.add_paragraph('Conforme consta nos autos, existem valores a serem restituídos à ré, tendo sido a ordem de transferência determinada por esse d. Juízo.',style='Paragraph')
    document.add_paragraph('Ocorre que, ainda que expedido ofício  ao gerente da instituição financeira depositante, para que fosse realizada transferência de valores em favor da seguradora Ré, não houve resposta do mesmo, com apresentação nos autos do respectivo comprovante.',style='Paragraph')
    document.add_paragraph('Assim, vem a Ré requerer a V. Exa., seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência realizada através de TED da quantia determinada em oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.',style='Paragraph')
    document.add_paragraph('Ademais, pugna-se que na requisição conste prazo para cumprimento da ordem judicial, sob pena de crime de desobediência, a fim de empregar plena efetividade e previsibilidade ao comando.',style='Paragraph')
    document.add_paragraph('Nestes Termos,',style='Paragraph')
    document.add_paragraph(' Pede Deferimento,',style='Paragraph')

    document.add_paragraph('MAMANGUAPE, 12 de maio de 2022.',style='Paragraph')

    document.add_paragraph('JOÃO BARBOSA',style='Paragraph')
    document.add_paragraph('OAB/PB 4246-A',style='Paragraph')


    document.add_paragraph('SUELIO MOREIRA TORRES',style='Paragraph')
    document.add_paragraph('15477 - OAB/PB',style='Paragraph')

    nome_do_arquivo='Peticao_DisponibilizacaoComprovanteTrans'+p_cod_cliente+'.docx'

    document.save('/home/ubuntu/documentos/'+nome_do_arquivo)
    return nome_do_arquivo
