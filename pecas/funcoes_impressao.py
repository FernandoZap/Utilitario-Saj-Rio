# -*- coding: utf-8 -*-

from docx import Document
from docx.shared import Length
import docx

import os
import datetime
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches,Cm
from . import funcoes_banco,funcoes_gerais
from .classes import Pasta

def pagina_inicial(doc,pasta_saj):
    #lista = funcoes_banco.f001_sql(pasta)

    #lista=funcoes_banco.view_pasta(pasta)
    #pasta_saj = Pasta(lista[0],lista[1],lista[2],lista[3],lista[4],lista[5],lista[6],lista[7],lista[8],lista[9],lista[10],lista[11],lista[12],lista[13],lista[14],lista[15],lista[16],lista[17],lista[18])
    dados_pasta = funcoes_banco.f002_sql(pasta_saj.pasta)

    font = doc.styles['Normal'].font
    font.name='Calibri'
    font.size=Pt(11)

    # margin doc
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(1)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    # paragrafo acima da logo jbaa
    paragraph1 = doc.add_paragraph()
    paragraph1_para = paragraph1.add_run(pasta_saj.cod_cliente +' - C1/ '+ pasta_saj.pasta+ '/ '+pasta_saj.cobertura)
    paragraph1.alignment=WD_ALIGN_PARAGRAPH.RIGHT
    paragraph1_para.font.size = Pt(11)

    # construcao da primeira pagina
    #imagem= doc.add_picture('C:/Users/Fernando/saj vps/img/logo_jbaa.png', width=Inches(1.25))
    imagem= doc.add_picture('/home/ubuntu/documentos/logo_jbaa.png', width=Inches(1.45))
    last_paragraph = doc.paragraphs[-1]
    last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph


    # paragrafo abaixo da logo jbaa
    paragraph2 = doc.add_paragraph()
    paragraph2_para = paragraph2.add_run('\n\n\nEXMO.SR.DR.JUIZ DE DIREITO ' + pasta_saj.juizo +' DA COMARCA DE ' + pasta_saj.comarca  +'/'+pasta_saj.uf+'\n\n\n')
    paragraph2_para.bold = True

    # nº processo
    paragraph3 = doc.add_paragraph()
    paragraph3_para = paragraph3.add_run('Processo: '+pasta_saj.nr_processo+'\n\n\n')
    paragraph3_para.alignment=WD_ALIGN_PARAGRAPH.LEFT
    paragraph3_para.bold = True

    paragraph4 = doc.add_paragraph()
    paragraph4_para = paragraph4
    paragraph4_para.add_run(pasta_saj.cliente).bold = True

    paragraph4_para.add_run(" já devidamente qualificadas nos autos do processo em epigrafe, por meio de seus advogados que esta subscreve, vem a presença de V. Excelencia nos autos da ")
    paragraph4_para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT ').bold = True 
    paragraph4_para.add_run('promovida por ')
    paragraph4_para.add_run(pasta_saj.autor).bold = True
    paragraph4_para.add_run(', OPOR EMBARGOS DE OMISSÃO, conforme passa a expor:')
    paragraph4_para.paragraph_format.first_line_indent = Inches(0.5)
    paragraph4_para.alignment=WD_ALIGN_PARAGRAPH.JUSTIFY


    # footer
    section = doc.sections[0]
    footer = section.footer
    footer_para = footer.paragraphs[0]
    para = footer_para.add_run(\
        ' \tRua São José, 90, 8º andar, Centro, Rio de Janeiro/RJ - CEP: 20010-020 \
        \n\twww.joaobarbosaadvass.com.br'\
    )
    para.font.size = Pt(9)
    para.font.name = 'Calibri'

    doc.add_page_break()



def pagina_de_encerramento(doc,pasta_saj):
    #lista = funcoes_banco.f001_sql(pasta)

    dados_pasta = funcoes_banco.f002_sql(pasta_saj.pasta)

    paragraph4=doc.add_paragraph('Neste Termos')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_after = Pt(1)
    paragraph4=doc.add_paragraph('Pede Deferimento')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER

    local = funcoes_gerais.local_e_data(dados_pasta['comarca'])
    paragraph4=doc.add_paragraph(local)
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run('\nJOÃO BARBOSA')
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4.paragraph_format.space_after = Pt(0)
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run(pasta_saj.oabjb)
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run('\n'+pasta_saj.conv_nome)
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4.paragraph_format.space_after = Pt(0)
    paragraph4_para.bold = True

    paragraph4=doc.add_paragraph()
    paragraph4_para = paragraph4.add_run(pasta_saj.conv_oab)
    paragraph4.alignment=WD_ALIGN_PARAGRAPH.CENTER
    paragraph4.paragraph_format.space_before = Pt(1)
    paragraph4_para.bold = True

'''
def pagina_inicial_ultrapetita(doc,pasta):
    lista = funcoes_banco.f001_sql(pasta)
    dados_pasta = funcoes_banco.f002_sql(pasta)

    pass



sec2= document.add_section(WD_SECTION.NEW_PAGE)
hd2 = se2.header
ft2 = sec2.footer

hd2.is_linked_to_previous = False
ft2.is_linked_to_previous = False
'''
