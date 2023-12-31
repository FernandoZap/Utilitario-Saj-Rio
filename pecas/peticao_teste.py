from docx import Document
import docx

from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
from docx.shared import RGBColor
from docx.shared import Inches


#document = Document()
#styles =  document.styles

#Style Paragraph
#p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
#p.font.name = "Calibri"
#p.font.size = Pt(11)

#Style Heading 2222
'''
h2 = styles.add_style("H2", WD_STYLE_TYPE.PARAGRAPH)
h2.base_style = styles["Heading 2"]
h2.font.name = "Calibri"
h2.font.size = Pt(13)
h2.font.color.rgb = RGBColor(78, 129, 189)
h2.font.bold = False

#Style Heading 3
h3 = styles.add_style("H3", WD_STYLE_TYPE.PARAGRAPH)
h3.base_style = styles["Heading 3"]
h3.font.name = "Calibri"
h3.font.size = Pt(12)
h3.font.color.rgb = RGBColor(78, 129, 189)
h3.font.bold = False
'''
def fun_peticao_02(p_pasta,p_autor,p_nr_processo,p_comarca,p_uf,p_cliente):
    doc = docx.Document()
    doc.add_heading('Peticao do documento', 0)
    para = doc.add_paragraph(
    '''GeeksforGeeks is a Computer Science portal for geeks.''')

    # Adding more content to paragraph and applying underline to them
    para.add_run(
    ''' It contains well written, well thought and well-explained ''').underline = True

    # Adding more content to paragraph
    para.add_run('''computer science and programming articles, quizzes etc.''')

    # Now save the document to a location

    doc.save('/home/ubuntu/documentos/teste2.docx')
    return 'teste2.docx'



def fun_peticao_01(p_pasta,p_cod_cliente,p_autor,p_nr_processo,p_comarca,p_uf,p_cliente,p_juizo,p_conv_nome,p_conv_oab):
    document = docx.Document()

    styles = document.styles
    p = styles.add_style("Paragraph", WD_STYLE_TYPE.PARAGRAPH)
    p.font.name = "Calibri"
    p.font.size = Pt(11)
    p.font.color.rgb=RGBColor(79, 129, 189)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)


    p2 = styles.add_style("Paragraph-2", WD_STYLE_TYPE.PARAGRAPH)
    p2.font.name = "Calibri"
    p2.font.size = Pt(11)
    p2.font.color.rgb=RGBColor(79, 129, 189)
    p2.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p2.paragraph_format.first_line_indent = Inches(0.5)
    p2.paragraph_format.left_indent = Inches(1.5)

    document.add_heading('EXMO SR. DR. JUIZ DE DIREITO DO(A) '+p_juizo+' DA COMARCA DE '+p_comarca+'/'+p_uf, 2)
    document.add_heading('Processo: '+p_nr_processo)
    para = document.add_paragraph(p_cliente+', previamente qualificada nos autos do processo em epígrafe, neste ato, representada por seus advogados que esta subscrevem, nos autos da ', style='Paragraph')
    under_and_bold = para.add_run('AÇÃO DE COBRANÇA DE SEGURO DPVAT')
    under_and_bold.underline=True
    under_and_bold.bold=True
    para.add_run(', que lhe promove ')
    para.add_run(p_autor).bold=True
    para.add_run(', em trâmite perante este Douto Juízo, vem respeitosamente, à presença de V. Exa., requerer o DESARQUIVAMENTO, a fim de viabilizar a DEVOLUÇÃO DOS HONORÁRIOS PERICIAIS PAGOS EM DUPLICIDADE (depósito judicial e ofício único de pagamento.')

    document.add_paragraph('Consoante se verifica nos autos e da documentação que segue em anexo, houve depósito a título de pagamento de honorários periciais, em cumprimento à intimação de fls., contudo, o processo foi relacionado para evento de mutirão de perícias, ocasião em que houve o pagamento da prova através de ofício único, restando, portanto, pagamento em duplicidade.',style='Paragraph')

    document.add_paragraph('Desta forma, com fulcro no art. 906, parágrafo único do CPC, requer a Ré que Vossa Excelência se digne determinar a expedição de OFÍCIO DE TRANSFERÊNCIA DIRETA no montante do valor depositado, com seus acréscimos legais, em favor da SEGURADORA LIDER DOS CONSÓRCIOS DO SEGURO DPVAT S.A., CNPJ/MF: 09.248.608/0001-04, autorizando ao Banco depositante a efetuar transferência direta na conta corrente nº 644000-2, Agência: 1912-7, BANCO DO BRASIL S.A', style='Paragraph')

    document.add_paragraph('Necessário esclarecer que a expedição da ordem de pagamento deverá ser nominal à SEGURADORA LÍDER DOS CONSÓRCIOS DO SEGURO DPVAT S/A, pois foi a empresa que custeou com o depósito como também é a gestora dos Consórcios do Seguro DPVAT nos termos do art. 5º, §3º, da Resolução CNSP de nº 154 , sendo a única e exclusiva beneficiária de reembolso da quantia disponível ao juízo.', style='Paragraph')

    document.add_paragraph('Reforçando o acima exposto, temos que as regras e os critérios para o DPVAT referentes aos sinistros ocorridos até 31 de dezembro de 2020 estão estabelecidas, também, na Resolução n.º 399 do CNSP de 29/12/2020.', style='Paragraph')

    document.add_paragraph('A referida Resolução prevê, no seu artigo 21, a competência da Seguradora Líder:',style='Paragraph')

    document.add_paragraph('Art. 21. A seguradora líder do Consórcio DPVAT será responsável pela gestão e operacionalização do seguro DPVAT referentes, exclusivamente, aos sinistros ocorridos até 31 de dezembro de 2020 (run-off), inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.',style='Paragraph-2')

    document.add_paragraph('Vejamos, agora, o art. 1º da Resolução 400 do CNSP de 29/12/2020:', style='Paragraph')

    document.add_paragraph('Art. 1º Ratificar que a Seguradora Líder do Consórcio do Seguro DPVAT S.A. será a responsável pela gestão e operacionalização do seguro DPVAT referentes, exclusivamente, aos sinistros ocorridos até 31 de dezembro de 2020, inclusive em relação às respectivas ações judiciais posteriormente ajuizadas.',style='Paragraph-2')

    document.add_paragraph('Requer ainda, seja determinado que o banco depositante junte aos autos o respectivo comprovante da transferência realizada através de TED da quantia expedida mediante oficio, possibilitando ao patrono da Ré realizar prestação de contas com maior clareza e transparência, informando o saldo líquido e a data exata da transferência realizada.',style='Paragraph')

    document.add_paragraph('Por fim, que seja observado exclusivamente o nome do advogado SUELIO MOREIRA TORRES ,15477/PB para efeito de intimações futuras, sob pena de nulidade das mesmas.',style='Paragraph')

    document.add_paragraph('Termos em que,',style='Paragraph')


    document.add_paragraph('Pede Juntada.',style='Paragraph')

    nome_do_arquivo = 'Peticao_teste_'+p_cod_cliente+'.docx'

    document.save('/home/ubuntu/documentos/'+nome_do_arquivo)

    return nome_do_arquivo
